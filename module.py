import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.shared import Pt, RGBColor, Inches
import os

class ReportGeneratorApp:
    def __init__(self, master):
        self.master = master
        master.title("Advanced Report Generator")

        self.projects = []

        tk.Button(master, text="Upload Program Files (.py/.r/.html)", command=self.upload_files).pack(pady=10)

        tk.Label(master, text="Select Language for Aim Statement:").pack()
        self.language_var = tk.StringVar(value="Python")
        ttk.Combobox(master, textvariable=self.language_var, values=["Python", "R", "HTML"]).pack()

        tk.Button(master, text="Generate Report", command=self.generate_report).pack(pady=10)

    def upload_files(self):
        file_paths = filedialog.askopenfilenames(filetypes=[("Program Files", "*.py *.r *.html")])
        if not file_paths:
            return

        for file_path in file_paths:
            code = self.read_file(file_path)
            lang = self.detect_language(file_path)
            title = self.generate_title_from_code(code, lang)
            self.projects.append({
                "heading": title,
                "code": code,
                "language": self.language_var.get()
            })

        messagebox.showinfo("Files Uploaded", f"{len(file_paths)} program(s) added successfully!")

    def read_file(self, file_path):
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    def detect_language(self, file_path):
        ext = os.path.splitext(file_path)[1].lower()
        return {"py": "Python", "r": "R", "html": "HTML"}.get(ext[1:], "Unknown")

    def generate_title_from_code(self, code, lang):
        if lang == "Python" and "def " in code:
            return code.split('def ')[1].split('(')[0].strip()
        elif lang == "R" and "<-" in code:
            return "Data Analysis Task"
        elif lang == "HTML" and "<title>" in code:
            return code.split('<title>')[1].split('</title>')[0].strip()
        return f"{lang} Program"

    def generate_algorithm(self, code, lang):
        steps = ["1. Start the program."]
        if lang == "Python":
            steps += ["2. Import necessary libraries.", "3. Define functions and logic."]
        elif lang == "R":
            steps += ["2. Load required R packages.", "3. Implement statistical or data operations."]
        elif lang == "HTML":
            steps += ["2. Create basic HTML structure.", "3. Add required page elements."]
        steps += ["4. Process/render outputs.", "5. End the program."]
        return "\n".join(steps)

    def add_bold_black_paragraph(self, doc, text, size=11, align=None, space_after=Pt(6)):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0, 0, 0)
        if align:
            p.alignment = align
        p.paragraph_format.space_after = space_after
        return p

    def add_black_paragraph(self, doc, text, size=11, space_after=Pt(6)):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.space_after = space_after
        return p

    def generate_report(self):
        if not self.projects:
            messagebox.showwarning("No Projects", "Please upload at least one project file.")
            return

        doc = Document()

        for exp_no, project in enumerate(self.projects, start=1):
            # Start a new section per experiment
            if exp_no > 1:
                doc.add_section(WD_SECTION_START.NEW_PAGE)

            # Top content
            self.add_bold_black_paragraph(doc, f"Exp. No: {exp_no}", size=12)
            self.add_bold_black_paragraph(doc, project['heading'].upper(), size=14, align=WD_PARAGRAPH_ALIGNMENT.CENTER)

            self.add_bold_black_paragraph(doc, "Aim:", size=12)
            aim = f"To write a program to {project['heading']} using {project['language']}."
            self.add_black_paragraph(doc, aim)

            self.add_bold_black_paragraph(doc, "Algorithm:", size=12)
            self.add_black_paragraph(doc, self.generate_algorithm(project['code'], project['language']))

            self.add_bold_black_paragraph(doc, "Program Code:", size=12)
            self.add_black_paragraph(doc, project['code'])

            # Add vertical spacing by inserting blank paragraphs
            for _ in range(12):  # Adjust count for margin space
                doc.add_paragraph()

            # Align "Result" to bottom by anchoring it after a space
            result_paragraph = self.add_bold_black_paragraph(doc, "Result:", size=12)
            result_paragraph.paragraph_format.space_before = Pt(120)

            self.add_black_paragraph(doc, "Thus the program has been successfully executed or created.")

        # Save the document
        doc.save("Final_Project_Report.docx")
        messagebox.showinfo("Report Generated", "Report saved as 'Final_Project_Report.docx'")

if __name__ == "__main__":
    root = tk.Tk()
    app = ReportGeneratorApp(root)
    root.mainloop()
